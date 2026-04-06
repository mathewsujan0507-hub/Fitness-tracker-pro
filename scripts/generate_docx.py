import sys
from pathlib import Path


def main() -> int:
    try:
        from docx import Document  # type: ignore
    except Exception:
        print("Missing dependency: python-docx")
        print("Install with: pip install -r requirements.txt")
        return 2

    root = Path(__file__).resolve().parents[1]
    docs_dir = root / "docs"
    out_dir = docs_dir
    out_dir.mkdir(parents=True, exist_ok=True)

    sources = [
        (docs_dir / "PROJECT_OVERVIEW.md", out_dir / "FitTrack-Pro_Detailed.docx"),
        (docs_dir / "PROJECT_OVERVIEW_SHORT.md", out_dir / "FitTrack-Pro_Short.docx"),
    ]

    for src, out in sources:
        if not src.exists():
            print(f"Missing source: {src}")
            return 3

        text = src.read_text(encoding="utf-8")
        doc = Document()

        for block in text.split("\n\n"):
            block = block.strip()
            if not block:
                continue

            if block.startswith("# "):
                doc.add_heading(block[2:].strip(), level=0)
                continue
            if block.startswith("## "):
                doc.add_heading(block[3:].strip(), level=1)
                continue
            if block.startswith("### "):
                doc.add_heading(block[4:].strip(), level=2)
                continue

            lines = block.splitlines()
            if all(l.strip().startswith(("- ", "* ")) for l in lines if l.strip()):
                for l in lines:
                    item = l.strip()[2:].strip()
                    if item:
                        doc.add_paragraph(item, style="List Bullet")
            else:
                doc.add_paragraph(block)

        doc.save(str(out))
        print(f"Wrote: {out}")

    return 0


if __name__ == "__main__":
    raise SystemExit(main())

